import re
from pathlib import Path
from typing import Any, Dict, List, Optional

from app.core.logging_config import logger


def _normalize_text(text: str) -> str:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    text = "\n".join(lines)

    while "\n\n\n" in text:
        text = text.replace("\n\n\n", "\n\n")

    return text.strip()


def _is_markdown_heading(line: str) -> bool:
    return bool(re.match(r"^\s{0,3}#{1,6}\s+\S+", line or ""))


def _strip_markdown_heading_prefix(line: str) -> str:
    return re.sub(r"^\s{0,3}#{1,6}\s+", "", line or "").strip()


def _markdown_heading_level(line: str) -> int:
    m = re.match(r"^\s{0,3}(#{1,6})\s+\S+", line or "")
    if not m:
        return 0
    return len(m.group(1))


def _is_likely_heading_line(line: str) -> bool:
    s = (line or "").strip()
    if not s:
        return False
    if len(s) > 80:
        return False
    if s.endswith(("：", ":")):
        return True
    if re.match(r"^(第[一二三四五六七八九十0-9]+[章节部分篇]|[0-9]+(\.[0-9]+)*\s+\S+)", s):
        return True
    if re.match(r"^(一|二|三|四|五|六|七|八|九|十)[、.．]\s*\S+", s):
        return True
    return False


def _extract_heading_from_chunk_text(text: str) -> str:
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    if not lines:
        return ""

    first = lines[0]
    m = re.match(r"^\[Heading\]\s*(.+)$", first, flags=re.IGNORECASE)
    if m:
        return m.group(1).strip()

    if _is_markdown_heading(first):
        return _strip_markdown_heading_prefix(first)

    if _is_likely_heading_line(first):
        return first

    return ""


def _should_keep_line_break(prev: str, cur: str) -> bool:
    prev = (prev or "").rstrip()
    cur = (cur or "").lstrip()
    if not prev or not cur:
        return True

    if _is_markdown_heading(cur) or _is_likely_heading_line(cur):
        return True
    if re.match(r"^[-*+]\s+\S+", cur):
        return True
    if re.match(r"^\d+[\.\)]\s+\S+", cur):
        return True
    if re.search(r"[。！？!?；;：:]$", prev):
        return True

    return False


def _merge_lines_smart(lines: List[str]) -> str:
    merged_parts: List[str] = []

    for raw in lines:
        line = (raw or "").strip()
        if not line:
            continue

        if not merged_parts:
            merged_parts.append(line)
            continue

        prev = merged_parts[-1]
        if _should_keep_line_break(prev, line):
            merged_parts.append(line)
        else:
            merged_parts[-1] = f"{prev} {line}".strip()

    return "\n".join(merged_parts).strip()


def _split_paragraphs(text: str) -> List[str]:
    text = _normalize_text(text)
    if not text:
        return []

    blocks = re.split(r"\n\s*\n+", text)
    paragraphs: List[str] = []

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        lines = [line.strip() for line in block.split("\n") if line.strip()]
        if not lines:
            continue

        merged = _merge_lines_smart(lines)
        if merged:
            paragraphs.append(merged)

    return paragraphs


def _hard_split_text(text: str, chunk_size: int, overlap: int) -> List[str]:
    text = _normalize_text(text)
    if not text:
        return []

    step = chunk_size - overlap
    if step <= 0:
        raise ValueError("chunk_size must be greater than overlap")

    chunks: List[str] = []
    start = 0

    while start < len(text):
        end = min(start + chunk_size, len(text))

        if end < len(text):
            window_start = max(start + int(chunk_size * 0.7), start)
            split_pos = -1
            for i in range(end, window_start, -1):
                ch = text[i - 1]
                if ch in "\n ，,。.!?！？；;：:、)]）】} ":
                    split_pos = i
                    break
            if split_pos != -1 and split_pos > start:
                end = split_pos

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        next_start = end - overlap
        if next_start <= start:
            next_start = start + step
        start = max(next_start, 0)

    return chunks


def _build_overlap_paragraphs(paragraphs: List[str], overlap: int) -> List[str]:
    if overlap <= 0 or not paragraphs:
        return []

    kept: List[str] = []
    total_len = 0

    for para in reversed(paragraphs):
        kept.append(para)
        total_len += len(para) + 2
        if total_len >= overlap:
            break

    kept.reverse()
    return kept


def _chunk_paragraph_list(
    paragraphs: List[str],
    chunk_size: int = 700,
    overlap: int = 100,
) -> List[str]:
    if not paragraphs:
        return []

    chunks: List[str] = []
    current_paragraphs: List[str] = []
    current_len = 0

    def flush_current() -> None:
        nonlocal current_paragraphs, current_len

        if not current_paragraphs:
            return

        chunk = "\n\n".join(current_paragraphs).strip()
        if chunk:
            chunks.append(chunk)

        overlap_paras = _build_overlap_paragraphs(current_paragraphs, overlap)
        current_paragraphs = overlap_paras[:]
        current_len = len("\n\n".join(current_paragraphs).strip()) if current_paragraphs else 0

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        if len(para) > chunk_size:
            flush_current()

            long_chunks = _hard_split_text(para, chunk_size=chunk_size, overlap=overlap)
            if not long_chunks:
                continue

            for lc in long_chunks[:-1]:
                if lc.strip():
                    chunks.append(lc.strip())

            last_piece = long_chunks[-1].strip()
            if last_piece:
                current_paragraphs = [last_piece]
                current_len = len(last_piece)
            else:
                current_paragraphs = []
                current_len = 0
            continue

        added_len = len(para) if current_len == 0 else current_len + 2 + len(para)

        if added_len <= chunk_size:
            current_paragraphs.append(para)
            current_len = added_len
        else:
            flush_current()
            current_paragraphs.append(para)
            current_len = len(para)

    if current_paragraphs:
        chunk = "\n\n".join(current_paragraphs).strip()
        if chunk:
            chunks.append(chunk)

    return chunks


def _chunk_markdown_sections(
    text: str,
    chunk_size: int = 700,
    overlap: int = 100,
) -> List[str]:
    text = _normalize_text(text)
    if not text:
        return []

    lines = text.split("\n")
    sections: List[Dict[str, Any]] = []
    current_heading: Optional[str] = None
    current_lines: List[str] = []

    def flush_section() -> None:
        nonlocal current_lines, current_heading
        body = "\n".join(current_lines).strip()
        if body or current_heading:
            sections.append(
                {
                    "heading": current_heading,
                    "body": body,
                }
            )
        current_lines = []

    for raw in lines:
        line = raw.rstrip()
        if _is_markdown_heading(line):
            flush_section()
            current_heading = _strip_markdown_heading_prefix(line)
        else:
            current_lines.append(line)

    flush_section()

    chunks: List[str] = []
    for section in sections:
        heading = (section["heading"] or "").strip()
        body = section["body"] or ""

        paragraphs = _split_paragraphs(body)
        if heading:
            paragraphs = [f"[Heading] {heading}"] + paragraphs

        section_chunks = _chunk_paragraph_list(
            paragraphs,
            chunk_size=chunk_size,
            overlap=overlap,
        )
        chunks.extend(section_chunks)

    return chunks


def _chunk_markdown_sections_with_meta(
    text: str,
    chunk_size: int = 700,
    overlap: int = 100,
) -> List[Dict[str, Any]]:
    text = _normalize_text(text)
    if not text:
        return []

    lines = text.split("\n")
    sections: List[Dict[str, Any]] = []
    current_heading: Optional[str] = None
    current_path: List[str] = []
    current_lines: List[str] = []
    heading_stack: List[tuple[int, str]] = []

    def flush_section() -> None:
        nonlocal current_lines, current_heading, current_path
        body = "\n".join(current_lines).strip()
        if body or current_heading:
            sections.append(
                {
                    "heading": current_heading,
                    "section_path": " > ".join(current_path) if current_path else "",
                    "body": body,
                }
            )
        current_lines = []

    for raw in lines:
        line = raw.rstrip()
        if _is_markdown_heading(line):
            flush_section()
            heading = _strip_markdown_heading_prefix(line)
            level = _markdown_heading_level(line)
            heading_stack = [item for item in heading_stack if item[0] < level]
            heading_stack.append((level, heading))
            current_heading = heading
            current_path = [item[1] for item in heading_stack]
        else:
            current_lines.append(line)

    flush_section()

    results: List[Dict[str, Any]] = []
    for section in sections:
        heading = (section["heading"] or "").strip()
        section_path = (section["section_path"] or "").strip()
        body = section["body"] or ""

        paragraphs = _split_paragraphs(body)
        if heading:
            paragraphs = [f"[Heading] {heading}"] + paragraphs

        section_chunks = _chunk_paragraph_list(
            paragraphs,
            chunk_size=chunk_size,
            overlap=overlap,
        )

        for chunk in section_chunks:
            chunk = chunk.strip()
            if not chunk:
                continue
            results.append(
                {
                    "text": chunk,
                    "page": 1,
                    "heading": heading or _extract_heading_from_chunk_text(chunk),
                    "section_path": section_path or heading,
                    "chunk_type": "section",
                    "source_type": "md",
                }
            )

    return results


def chunk_text(
    text: str,
    chunk_size: int = 700,
    overlap: int = 100,
    source_type: str = "plain",
) -> List[str]:
    text = _normalize_text(text)
    if not text:
        return []

    if chunk_size <= overlap:
        raise ValueError("chunk_size must be greater than overlap")

    source_type = (source_type or "plain").lower()

    if source_type == "md":
        chunks = _chunk_markdown_sections(text, chunk_size=chunk_size, overlap=overlap)
        logger.info(
            "chunk_text markdown done text_len=%s chunk_size=%s overlap=%s chunk_count=%s",
            len(text),
            chunk_size,
            overlap,
            len(chunks),
        )
        return chunks

    paragraphs = _split_paragraphs(text)
    if not paragraphs:
        chunks = _hard_split_text(text, chunk_size=chunk_size, overlap=overlap)
        logger.info(
            "chunk_text fallback_hard_split text_len=%s chunk_size=%s overlap=%s chunk_count=%s",
            len(text),
            chunk_size,
            overlap,
            len(chunks),
        )
        return chunks

    chunks = _chunk_paragraph_list(paragraphs, chunk_size=chunk_size, overlap=overlap)
    logger.info(
        "chunk_text done text_len=%s paragraph_count=%s chunk_size=%s overlap=%s chunk_count=%s source_type=%s",
        len(text),
        len(paragraphs),
        chunk_size,
        overlap,
        len(chunks),
        source_type,
    )
    return chunks


def _read_pdf_pages(path: Path) -> List[Dict[str, Any]]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: List[Dict[str, Any]] = []

    for page_num, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception as e:
            logger.warning(
                "read_pdf_pages page extract failed path=%s page=%s error=%s",
                str(path),
                page_num,
                str(e),
            )
            page_text = ""

        page_text = _normalize_text(page_text)
        if page_text:
            pages.append({"page": page_num, "text": page_text})

    logger.info(
        "_read_pdf_pages done path=%s page_count=%s extracted_page_count=%s",
        str(path),
        len(reader.pages),
        len(pages),
    )
    return pages


def _read_pdf_file(path: Path) -> str:
    pages = _read_pdf_pages(path)
    return "\n\n".join(item["text"] for item in pages).strip()


def _read_docx_file(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: List[str] = []

    for para in doc.paragraphs:
        text = _normalize_text(para.text or "")
        if text:
            parts.append(text)

    for table in doc.tables:
        row_texts: List[str] = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = _normalize_text(cell.text or "")
                if cell_text:
                    cells.append(cell_text)
            if cells:
                row_texts.append(" | ".join(cells))
        if row_texts:
            parts.append("\n".join(row_texts))

    text = "\n\n".join(parts).strip()
    logger.info(
        "_read_docx_file done path=%s paragraph_count=%s table_count=%s text_len=%s",
        str(path),
        len(doc.paragraphs),
        len(doc.tables),
        len(text),
    )
    return text


def read_text_file(source_path: str) -> str:
    path = Path(source_path)
    logger.info("read_text_file source_path=%s", source_path)

    if not path.exists():
        raise FileNotFoundError(f"file not found: {source_path}")

    ext = path.suffix.lower()

    if ext in [".txt", ".md"]:
        text = path.read_text(encoding="utf-8")
    elif ext == ".pdf":
        text = _read_pdf_file(path)
    elif ext == ".docx":
        text = _read_docx_file(path)
    else:
        raise ValueError(f"unsupported file type: {ext}")

    text = _normalize_text(text)

    logger.info(
        "read_text_file done source_path=%s ext=%s text_len=%s",
        source_path,
        ext,
        len(text),
    )
    return text


def _chunk_page_text(
    text: str,
    page: int | None,
    chunk_size: int = 700,
    overlap: int = 100,
) -> List[Dict[str, Any]]:
    paragraph_chunks = chunk_text(
        text,
        chunk_size=chunk_size,
        overlap=overlap,
        source_type="pdf",
    )
    results: List[Dict[str, Any]] = []
    for chunk in paragraph_chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        heading = _extract_heading_from_chunk_text(chunk)
        results.append(
            {
                "text": chunk,
                "page": page,
                "heading": heading,
                "section_path": heading,
                "chunk_type": "page",
                "source_type": "pdf",
            }
        )
    return results


def read_chunks_with_meta(
    source_path: str,
    chunk_size: int = 700,
    overlap: int = 100,
) -> List[Dict[str, Any]]:
    path = Path(source_path)
    logger.info("read_chunks_with_meta source_path=%s", source_path)

    if not path.exists():
        raise FileNotFoundError(f"file not found: {source_path}")

    ext = path.suffix.lower()

    if ext == ".pdf":
        results: List[Dict[str, Any]] = []
        for item in _read_pdf_pages(path):
            results.extend(
                _chunk_page_text(
                    text=item["text"],
                    page=item["page"],
                    chunk_size=chunk_size,
                    overlap=overlap,
                )
            )

        logger.info(
            "read_chunks_with_meta pdf done source_path=%s chunk_count=%s",
            source_path,
            len(results),
        )
        return results

    text = read_text_file(source_path)
    if ext == ".md":
        results = _chunk_markdown_sections_with_meta(
            text,
            chunk_size=chunk_size,
            overlap=overlap,
        )
    else:
        source_type = ext.lstrip(".") or "plain"
        text_chunks = chunk_text(
            text,
            chunk_size=chunk_size,
            overlap=overlap,
            source_type=source_type,
        )
        results = []
        for chunk in text_chunks:
            chunk = chunk.strip()
            if not chunk:
                continue
            heading = _extract_heading_from_chunk_text(chunk)
            results.append(
                {
                    "text": chunk,
                    "page": 1,
                    "heading": heading,
                    "section_path": heading,
                    "chunk_type": "paragraph",
                    "source_type": source_type,
                }
            )

    logger.info(
        "read_chunks_with_meta non_pdf done source_path=%s chunk_count=%s ext=%s",
        source_path,
        len(results),
        ext,
    )
    return results
